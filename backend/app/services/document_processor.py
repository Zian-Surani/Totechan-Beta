import os
import re
import uuid
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import structlog
from io import BytesIO

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import chardet

from app.config.settings import settings
from app.utils.exceptions import DocumentProcessingError
from app.models.document_schemas import DocumentType

logger = structlog.get_logger()


class DocumentProcessor:
    """Service for processing various document types and extracting text"""

    def __init__(self):
        self.chunk_size = settings.default_chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.supported_types = {
            DocumentType.PDF: self._process_pdf,
            DocumentType.DOCX: self._process_docx,
            DocumentType.HTML: self._process_html,
            DocumentType.TXT: self._process_txt,
        }

    async def process_document(
        self,
        file_path: str,
        file_type: DocumentType,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Process document and extract text with metadata"""
        try:
            if file_type not in self.supported_types:
                raise DocumentProcessingError(
                    f"Unsupported document type: {file_type}",
                    details={"supported_types": list(self.supported_types.keys())}
                )

            # Check if file exists
            if not os.path.exists(file_path):
                raise DocumentProcessingError(f"File not found: {file_path}")

            # Process document based on type
            processor_func = self.supported_types[file_type]
            result = await processor_func(file_path, metadata or {})

            # Add document processing metadata
            result.update({
                "file_type": file_type.value,
                "processed_at": str(uuid.uuid4()),  # Unique processing ID
                "chunk_count": len(result["chunks"])
            })

            logger.info(
                "document_processed",
                file_path=file_path,
                file_type=file_type.value,
                chunk_count=result["chunk_count"],
                total_chars=result.get("total_chars", 0)
            )

            return result

        except Exception as e:
            logger.error(
                "document_processing_error",
                file_path=file_path,
                file_type=file_type.value,
                error=str(e)
            )
            raise DocumentProcessingError(
                f"Failed to process document: {str(e)}",
                details={"file_path": file_path, "file_type": file_type.value}
            )

    async def _process_pdf(
        self,
        file_path: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process PDF document"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract metadata
                pdf_info = pdf_reader.metadata
                document_metadata = {
                    "title": pdf_info.title if pdf_info and pdf_info.title else "",
                    "author": pdf_info.author if pdf_info and pdf_info.author else "",
                    "creator": pdf_info.creator if pdf_info else "",
                    "producer": pdf_info.producer if pdf_info else "",
                    "creation_date": str(pdf_info.creation_date) if pdf_info and pdf_info.creation_date else "",
                    "page_count": len(pdf_reader.pages),
                    **metadata
                }

                # Extract text from all pages
                text_content = ""
                pages_text = []

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += page_text + "\n"
                            pages_text.append({
                                "page_number": page_num + 1,
                                "text": page_text.strip(),
                                "char_count": len(page_text)
                            })
                    except Exception as e:
                        logger.warning(
                            "pdf_page_extraction_error",
                            page_num=page_num,
                            error=str(e)
                        )
                        continue

                # Create chunks
                chunks = self._create_chunks(
                    text_content,
                    document_metadata,
                    {"pages": pages_text}
                )

                return {
                    "metadata": document_metadata,
                    "text": text_content,
                    "chunks": chunks,
                    "total_chars": len(text_content),
                    "pages": pages_text
                }

        except Exception as e:
            raise DocumentProcessingError(
                f"PDF processing failed: {str(e)}",
                details={"file_path": file_path}
            )

    async def _process_docx(
        self,
        file_path: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process DOCX document"""
        try:
            doc = Document(file_path)

            # Extract metadata
            core_props = doc.core_properties
            document_metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "keywords": core_props.keywords or "",
                **metadata
            }

            # Extract text from paragraphs
            paragraphs_text = []
            text_content = ""

            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text_content += para_text + "\n"
                    paragraphs_text.append({
                        "paragraph": len(paragraphs_text) + 1,
                        "text": para_text,
                        "style": para.style.name if para.style else "Normal"
                    })

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_content.append(cell_text)
                    table_content.append(" | ".join(row_content))

                if table_content:
                    table_text_str = "\n".join(table_content)
                    text_content += "\n" + table_text_str + "\n"
                    tables_text.append({
                        "table": len(tables_text) + 1,
                        "text": table_text_str,
                        "rows": len(table_content)
                    })

            # Create chunks
            chunks = self._create_chunks(
                text_content,
                document_metadata,
                {
                    "paragraphs": paragraphs_text,
                    "tables": tables_text
                }
            )

            return {
                "metadata": document_metadata,
                "text": text_content,
                "chunks": chunks,
                "total_chars": len(text_content),
                "paragraphs": paragraphs_text,
                "tables": tables_text
            }

        except Exception as e:
            raise DocumentProcessingError(
                f"DOCX processing failed: {str(e)}",
                details={"file_path": file_path}
            )

    async def _process_html(
        self,
        file_path: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process HTML document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Extract metadata
            title_tag = soup.find('title')
            meta_tags = soup.find_all('meta')

            document_metadata = {
                "title": title_tag.get_text().strip() if title_tag else "",
                "description": "",
                "keywords": "",
                "author": "",
                **metadata
            }

            # Extract meta information
            for meta in meta_tags:
                name = meta.get('name', '').lower()
                content = meta.get('content', '')

                if name == 'description':
                    document_metadata["description"] = content
                elif name == 'keywords':
                    document_metadata["keywords"] = content
                elif name == 'author':
                    document_metadata["author"] = content

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text content
            text_content = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = '\n'.join(chunk for chunk in chunks if chunk)

            # Extract structured content
            headings = []
            for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                headings.append({
                    "level": int(h.name[1]),
                    "text": h.get_text().strip()
                })

            # Create chunks
            chunks = self._create_chunks(
                text_content,
                document_metadata,
                {
                    "headings": headings,
                    "url": metadata.get("url", "")
                }
            )

            return {
                "metadata": document_metadata,
                "text": text_content,
                "chunks": chunks,
                "total_chars": len(text_content),
                "headings": headings
            }

        except Exception as e:
            raise DocumentProcessingError(
                f"HTML processing failed: {str(e)}",
                details={"file_path": file_path}
            )

    async def _process_txt(
        self,
        file_path: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process TXT document"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'

            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding) as file:
                text_content = file.read()

            # Basic metadata
            document_metadata = {
                "encoding": encoding,
                "line_count": len(text_content.splitlines()),
                **metadata
            }

            # Clean up text
            text_content = text_content.strip()

            # Create chunks
            chunks = self._create_chunks(
                text_content,
                document_metadata,
                {"encoding": encoding}
            )

            return {
                "metadata": document_metadata,
                "text": text_content,
                "chunks": chunks,
                "total_chars": len(text_content)
            }

        except Exception as e:
            raise DocumentProcessingError(
                f"TXT processing failed: {str(e)}",
                details={"file_path": file_path, "encoding": encoding if 'encoding' in locals() else "unknown"}
            )

    def _create_chunks(
        self,
        text: str,
        metadata: Dict[str, Any],
        extra_info: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """Create text chunks with overlap"""
        try:
            # Clean text first
            text = re.sub(r'\s+', ' ', text).strip()

            if not text:
                return []

            chunks = []
            start = 0
            chunk_index = 0

            while start < len(text):
                # Calculate chunk end
                end = start + self.chunk_size

                # Try to break at word boundary
                if end < len(text):
                    # Find the last space before the end
                    last_space = text.rfind(' ', start, end)
                    if last_space != -1:
                        end = last_space

                # Extract chunk text
                chunk_text = text[start:end].strip()

                if chunk_text:
                    # Create chunk metadata
                    chunk_metadata = {
                        "doc_id": metadata.get("doc_id", ""),
                        "source_filename": metadata.get("filename", ""),
                        "file_type": metadata.get("file_type", ""),
                        "chunk_index": chunk_index,
                        "start_char": start,
                        "end_char": end,
                        "char_count": len(chunk_text),
                        "created_at": metadata.get("created_at", ""),
                        **(extra_info or {})
                    }

                    # Add page information for PDFs
                    if "pages" in extra_info:
                        chunk_metadata["page_number"] = self._get_page_for_chunk(
                            start, end, extra_info["pages"]
                        )

                    chunks.append({
                        "text": chunk_text,
                        "metadata": chunk_metadata
                    })

                    chunk_index += 1

                # Calculate next start with overlap
                start = max(start + 1, end - self.chunk_overlap)

            return chunks

        except Exception as e:
            logger.error("chunk_creation_error", error=str(e))
            # Fallback: return single chunk
            return [{
                "text": text,
                "metadata": {
                    "doc_id": metadata.get("doc_id", ""),
                    "source_filename": metadata.get("filename", ""),
                    "file_type": metadata.get("file_type", ""),
                    "chunk_index": 0,
                    "start_char": 0,
                    "end_char": len(text),
                    "char_count": len(text),
                    "created_at": metadata.get("created_at", ""),
                    **(extra_info or {})
                }
            }]

    def _get_page_for_chunk(
        self,
        start: int,
        end: int,
        pages: List[Dict[str, Any]]
    ) -> Optional[int]:
        """Determine which page a chunk belongs to based on character positions"""
        if not pages:
            return None

        # Simple heuristic: assume pages have roughly equal character distribution
        # This is a basic implementation - more sophisticated approaches exist
        total_chars = sum(page.get("char_count", 0) for page in pages)
        if total_chars == 0:
            return 1

        # Find approximate page based on character position
        position_ratio = start / total_chars if total_chars > 0 else 0
        estimated_page = int(position_ratio * len(pages)) + 1

        # Ensure within bounds
        return max(1, min(estimated_page, len(pages)))

    async def estimate_processing_time(
        self,
        file_size_bytes: int,
        file_type: DocumentType
    ) -> float:
        """Estimate processing time in seconds"""
        # Simple heuristic based on file size and type
        base_times = {
            DocumentType.PDF: 2.0,   # PDFs are slower to process
            DocumentType.DOCX: 1.5,  # DOCX medium speed
            DocumentType.HTML: 1.0,  # HTML relatively fast
            DocumentType.TXT: 0.5,   # TXT fastest
        }

        base_time = base_times.get(file_type, 1.0)

        # Scale by file size (1MB as baseline)
        size_mb = file_size_bytes / (1024 * 1024)
        estimated_time = base_time * (1 + size_mb)

        return max(1.0, estimated_time)  # Minimum 1 second